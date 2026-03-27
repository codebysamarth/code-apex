"""
Document Parser
===============
Parses various document formats into plain text for BRD extraction.

Supported formats:
- Plain text (.txt)
- PDF documents (.pdf)
- Word documents (.docx)
- Meeting transcripts (.vtt, .srt)
- Markdown (.md)

Dependencies:
- pdfplumber (PDF parsing)
- python-docx (Word parsing)
"""

import os
import re
from typing import Optional
from dataclasses import dataclass


@dataclass
class ParsedDocument:
    """Result of document parsing."""
    text: str
    source_type: str
    filename: str
    page_count: int = 1
    word_count: int = 0
    metadata: dict = None
    
    def __post_init__(self):
        if self.metadata is None:
            self.metadata = {}
        self.word_count = len(self.text.split())


class DocumentParser:
    """Parse various document formats into plain text."""
    
    SUPPORTED_EXTENSIONS = {
        '.txt': 'text',
        '.md': 'markdown',
        '.pdf': 'pdf',
        '.docx': 'word',
        '.doc': 'word',
        '.vtt': 'transcript',
        '.srt': 'transcript',
    }
    
    def __init__(self):
        self._pdf_available = False
        self._docx_available = False
        self._check_dependencies()
    
    def _check_dependencies(self):
        """Check which optional dependencies are available."""
        try:
            import pdfplumber
            self._pdf_available = True
        except ImportError:
            print("[DocumentParser] pdfplumber not installed - PDF support disabled")
        
        try:
            import docx
            self._docx_available = True
        except ImportError:
            print("[DocumentParser] python-docx not installed - Word support disabled")
    
    def parse(self, content: bytes, filename: str) -> ParsedDocument:
        """
        Parse document content based on file extension.
        
        Args:
            content: Raw file bytes
            filename: Original filename (used to determine format)
            
        Returns:
            ParsedDocument with extracted text and metadata
        """
        ext = os.path.splitext(filename.lower())[1]
        
        if ext not in self.SUPPORTED_EXTENSIONS:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {list(self.SUPPORTED_EXTENSIONS.keys())}")
        
        source_type = self.SUPPORTED_EXTENSIONS[ext]
        
        if ext == '.txt':
            return self._parse_text(content, filename)
        elif ext == '.md':
            return self._parse_markdown(content, filename)
        elif ext == '.pdf':
            return self._parse_pdf(content, filename)
        elif ext in ('.docx', '.doc'):
            return self._parse_docx(content, filename)
        elif ext in ('.vtt', '.srt'):
            return self._parse_transcript(content, filename, ext)
        else:
            return self._parse_text(content, filename)
    
    def parse_string(self, text: str, source_type: str = "text") -> ParsedDocument:
        """Parse a plain text string (for API compatibility)."""
        return ParsedDocument(
            text=text.strip(),
            source_type=source_type,
            filename="input.txt",
            page_count=1,
        )
    
    def _parse_text(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse plain text file."""
        # Try different encodings
        for encoding in ['utf-8', 'utf-16', 'latin-1', 'cp1252']:
            try:
                text = content.decode(encoding)
                break
            except UnicodeDecodeError:
                continue
        else:
            text = content.decode('utf-8', errors='replace')
        
        return ParsedDocument(
            text=text.strip(),
            source_type='text',
            filename=filename,
            page_count=1,
        )
    
    def _parse_markdown(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse Markdown file, converting to plain text."""
        text = content.decode('utf-8', errors='replace')
        
        # Remove markdown formatting but keep structure
        # Remove code blocks
        text = re.sub(r'```[\s\S]*?```', '', text)
        text = re.sub(r'`[^`]+`', '', text)
        
        # Convert headers to plain text with newlines
        text = re.sub(r'^#+\s*', '', text, flags=re.MULTILINE)
        
        # Remove links but keep text
        text = re.sub(r'\[([^\]]+)\]\([^)]+\)', r'\1', text)
        
        # Remove bold/italic markers
        text = re.sub(r'\*\*([^*]+)\*\*', r'\1', text)
        text = re.sub(r'\*([^*]+)\*', r'\1', text)
        text = re.sub(r'__([^_]+)__', r'\1', text)
        text = re.sub(r'_([^_]+)_', r'\1', text)
        
        # Remove images
        text = re.sub(r'!\[([^\]]*)\]\([^)]+\)', '', text)
        
        # Clean up multiple newlines
        text = re.sub(r'\n{3,}', '\n\n', text)
        
        return ParsedDocument(
            text=text.strip(),
            source_type='markdown',
            filename=filename,
            page_count=1,
        )
    
    def _parse_pdf(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse PDF document using pdfplumber."""
        if not self._pdf_available:
            raise ImportError("pdfplumber not installed. Run: pip install pdfplumber")
        
        import pdfplumber
        import io
        
        text_parts = []
        page_count = 0
        metadata = {}
        
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            page_count = len(pdf.pages)
            metadata = pdf.metadata or {}
            
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        
        text = '\n\n'.join(text_parts)
        
        # Clean up common PDF artifacts
        text = re.sub(r'\s+', ' ', text)  # Normalize whitespace
        text = re.sub(r'(\w)-\s+(\w)', r'\1\2', text)  # Fix hyphenation
        
        return ParsedDocument(
            text=text.strip(),
            source_type='pdf',
            filename=filename,
            page_count=page_count,
            metadata={
                'title': metadata.get('Title', ''),
                'author': metadata.get('Author', ''),
                'created': str(metadata.get('CreationDate', '')),
            }
        )
    
    def _parse_docx(self, content: bytes, filename: str) -> ParsedDocument:
        """Parse Word document using python-docx."""
        if not self._docx_available:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        
        from docx import Document
        import io
        
        doc = Document(io.BytesIO(content))
        
        text_parts = []
        
        # Extract paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(row_text)
        
        text = '\n\n'.join(text_parts)
        
        # Get metadata
        core_props = doc.core_properties
        metadata = {
            'title': core_props.title or '',
            'author': core_props.author or '',
            'created': str(core_props.created) if core_props.created else '',
            'modified': str(core_props.modified) if core_props.modified else '',
        }
        
        return ParsedDocument(
            text=text.strip(),
            source_type='word',
            filename=filename,
            page_count=1,  # DOCX doesn't have page concept in API
            metadata=metadata,
        )
    
    def _parse_transcript(self, content: bytes, filename: str, ext: str) -> ParsedDocument:
        """Parse meeting transcript (VTT or SRT format)."""
        text = content.decode('utf-8', errors='replace')
        
        if ext == '.vtt':
            return self._parse_vtt(text, filename)
        else:
            return self._parse_srt(text, filename)
    
    def _parse_vtt(self, text: str, filename: str) -> ParsedDocument:
        """Parse WebVTT subtitle format."""
        lines = text.split('\n')
        transcript_parts = []
        current_speaker = None
        
        # Skip header
        i = 0
        while i < len(lines) and not re.match(r'\d{2}:\d{2}', lines[i]):
            i += 1
        
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip timestamp lines
            if re.match(r'\d{2}:\d{2}:\d{2}', line):
                i += 1
                continue
            
            # Skip cue identifiers
            if re.match(r'^\d+$', line) or not line:
                i += 1
                continue
            
            # Check for speaker identification
            speaker_match = re.match(r'^<v\s+([^>]+)>(.*)$', line)
            if speaker_match:
                speaker = speaker_match.group(1)
                speech = speaker_match.group(2).strip()
                if speaker != current_speaker:
                    current_speaker = speaker
                    transcript_parts.append(f"\n{speaker}: {speech}")
                else:
                    transcript_parts.append(speech)
            else:
                # Plain text without speaker tag
                clean_line = re.sub(r'<[^>]+>', '', line)  # Remove HTML tags
                if clean_line.strip():
                    transcript_parts.append(clean_line)
            
            i += 1
        
        text = ' '.join(transcript_parts)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return ParsedDocument(
            text=text,
            source_type='transcript',
            filename=filename,
            page_count=1,
            metadata={'format': 'vtt'}
        )
    
    def _parse_srt(self, text: str, filename: str) -> ParsedDocument:
        """Parse SubRip subtitle format."""
        # Remove sequence numbers and timestamps
        lines = text.split('\n')
        transcript_parts = []
        
        i = 0
        while i < len(lines):
            line = lines[i].strip()
            
            # Skip sequence numbers
            if re.match(r'^\d+$', line):
                i += 1
                continue
            
            # Skip timestamp lines
            if re.match(r'\d{2}:\d{2}:\d{2},\d{3}\s*-->', line):
                i += 1
                continue
            
            # Skip empty lines
            if not line:
                i += 1
                continue
            
            # This is actual text
            clean_line = re.sub(r'<[^>]+>', '', line)  # Remove HTML tags
            if clean_line.strip():
                transcript_parts.append(clean_line)
            
            i += 1
        
        text = ' '.join(transcript_parts)
        text = re.sub(r'\s+', ' ', text).strip()
        
        return ParsedDocument(
            text=text,
            source_type='transcript',
            filename=filename,
            page_count=1,
            metadata={'format': 'srt'}
        )
    
    def get_supported_formats(self) -> dict:
        """Return dict of supported formats and their availability."""
        return {
            '.txt': True,
            '.md': True,
            '.pdf': self._pdf_available,
            '.docx': self._docx_available,
            '.doc': self._docx_available,
            '.vtt': True,
            '.srt': True,
        }


# Singleton instance
_parser = None

def get_parser() -> DocumentParser:
    """Get singleton document parser instance."""
    global _parser
    if _parser is None:
        _parser = DocumentParser()
    return _parser


if __name__ == "__main__":
    # Test the parser
    parser = DocumentParser()
    
    print("Supported formats:")
    for fmt, available in parser.get_supported_formats().items():
        status = "✓" if available else "✗ (missing dependency)"
        print(f"  {fmt}: {status}")
    
    # Test plain text parsing
    sample = b"The system must support user authentication.\nPayments should process within 3 seconds."
    result = parser.parse(sample, "test.txt")
    print(f"\nParsed text document:")
    print(f"  Source type: {result.source_type}")
    print(f"  Word count: {result.word_count}")
    print(f"  Preview: {result.text[:100]}...")
