"""
Word Export
===========
Export BRD to Microsoft Word format (.docx).

Uses python-docx for document generation with professional formatting.
"""

import io
from datetime import datetime
from typing import Optional


def _check_docx():
    """Check if python-docx is available."""
    try:
        from docx import Document
        from docx.shared import Inches, Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.table import WD_TABLE_ALIGNMENT
        return True
    except ImportError:
        return False


def export_to_word(brd_data: dict, output_path: Optional[str] = None) -> bytes:
    """
    Export BRD data to Word format (.docx).
    
    Args:
        brd_data: The BRD extraction result dictionary
        output_path: Optional file path to save document (if None, returns bytes)
        
    Returns:
        DOCX content as bytes
    """
    if not _check_docx():
        raise ImportError("python-docx not installed. Run: pip install python-docx")
    
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    
    doc = Document()
    
    # Document properties
    doc.core_properties.title = brd_data.get('project_name', 'Business Requirements Document')
    doc.core_properties.author = 'BRD Agent'
    
    # Title
    title = doc.add_heading(brd_data.get('project_name', 'Business Requirements Document'), 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Subtitle with date
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = subtitle.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    run.font.color.rgb = RGBColor(113, 128, 150)
    run.font.size = Pt(11)
    
    doc.add_paragraph()  # Spacer
    
    # Executive Summary
    doc.add_heading('Executive Summary', level=1)
    
    fr_count = len(brd_data.get('functional_reqs', []))
    nfr_count = len(brd_data.get('nfrs', []))
    stakeholder_count = len(brd_data.get('stakeholders', []))
    decision_count = len(brd_data.get('decisions', []))
    completeness = brd_data.get('completeness_score', 0) * 100
    
    summary = doc.add_paragraph()
    summary.add_run('Document Overview: ').bold = True
    summary.add_run(
        f"This BRD contains {fr_count} functional requirements, {nfr_count} non-functional requirements, "
        f"{stakeholder_count} stakeholders, and {decision_count} key decisions."
    )
    
    score_para = doc.add_paragraph()
    score_para.add_run('Completeness Score: ').bold = True
    score_para.add_run(f"{completeness:.0f}%")
    
    # Functional Requirements
    frs = brd_data.get('functional_reqs', [])
    if frs:
        doc.add_heading('Functional Requirements', level=1)
        
        table = doc.add_table(rows=1, cols=5)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # Header row
        header_cells = table.rows[0].cells
        headers = ['ID', 'Requirement', 'Priority', 'MoSCoW', 'Confidence']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            # Make header bold and add background
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(10)
        
        # Data rows
        for req in frs:
            row_cells = table.add_row().cells
            row_cells[0].text = req.get('id', '')
            row_cells[1].text = req.get('text', '')[:150]
            row_cells[2].text = str(req.get('priority', ''))
            row_cells[3].text = req.get('moscow', '')
            row_cells[4].text = f"{req.get('confidence', 0):.0%}"
        
        # Add reasoning section if available
        doc.add_paragraph()
        has_reasoning = any(req.get('reasoning') for req in frs)
        if has_reasoning:
            reasoning_heading = doc.add_paragraph()
            reasoning_heading.add_run('Extraction Reasoning:').bold = True
            for req in frs:
                if req.get('reasoning'):
                    p = doc.add_paragraph(style='List Bullet')
                    p.add_run(f"{req.get('id', '')}: ").bold = True
                    p.add_run(req.get('reasoning', ''))
    
    # Non-Functional Requirements
    nfrs = brd_data.get('nfrs', [])
    if nfrs:
        doc.add_heading('Non-Functional Requirements', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        headers = ['ID', 'Requirement', 'Category', 'Confidence']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for req in nfrs:
            row_cells = table.add_row().cells
            row_cells[0].text = req.get('id', '')
            row_cells[1].text = req.get('text', '')[:180]
            row_cells[2].text = req.get('category', '')
            row_cells[3].text = f"{req.get('confidence', 0):.0%}"
    
    # Stakeholders
    stakeholders = brd_data.get('stakeholders', [])
    if stakeholders:
        doc.add_heading('Stakeholders', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        headers = ['Name', 'Role', 'Influence Score', 'Mentions']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for sh in stakeholders:
            row_cells = table.add_row().cells
            row_cells[0].text = sh.get('name', '')
            row_cells[1].text = sh.get('role', '')
            row_cells[2].text = f"{sh.get('influence_score', 0):.0%}"
            row_cells[3].text = str(sh.get('mentioned_count', 0))
    
    # Decisions
    decisions = brd_data.get('decisions', [])
    if decisions:
        doc.add_heading('Key Decisions', level=1)
        
        for dec in decisions:
            p = doc.add_paragraph()
            p.add_run(f"{dec.get('id', '')}: ").bold = True
            p.add_run(dec.get('text', ''))
            
            if dec.get('rationale'):
                rationale = doc.add_paragraph()
                rationale.add_run('Rationale: ').italic = True
                rationale.add_run(dec.get('rationale', '')).italic = True
            
            if dec.get('date_mentioned'):
                date_p = doc.add_paragraph()
                date_p.add_run(f"Date: {dec.get('date_mentioned')}")
    
    # Scope
    scope = brd_data.get('scope', {})
    if any(scope.values()):
        doc.add_heading('Scope', level=1)
        
        if scope.get('in_scope'):
            in_scope_heading = doc.add_paragraph()
            in_scope_heading.add_run('In Scope:').bold = True
            for item in scope['in_scope']:
                doc.add_paragraph(item, style='List Bullet')
        
        if scope.get('out_of_scope'):
            out_scope_heading = doc.add_paragraph()
            out_scope_heading.add_run('Out of Scope:').bold = True
            for item in scope['out_of_scope']:
                doc.add_paragraph(item, style='List Bullet')
        
        if scope.get('assumptions'):
            assumptions_heading = doc.add_paragraph()
            assumptions_heading.add_run('Assumptions:').bold = True
            for item in scope['assumptions']:
                doc.add_paragraph(item, style='List Bullet')
    
    # Timeline
    timeline = brd_data.get('timeline', [])
    if timeline:
        doc.add_heading('Timeline & Milestones', level=1)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        header_cells = table.rows[0].cells
        headers = ['Milestone', 'Date', 'Type', 'Urgency']
        for i, header in enumerate(headers):
            header_cells[i].text = header
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
        
        for ms in timeline:
            row_cells = table.add_row().cells
            row_cells[0].text = ms.get('milestone', '')[:80]
            row_cells[1].text = ms.get('date', '')
            row_cells[2].text = ms.get('type', '')
            row_cells[3].text = ms.get('urgency', '')
    
    # Gaps
    gaps = brd_data.get('gaps', [])
    if gaps:
        doc.add_heading('Identified Gaps', level=1)
        
        for gap in gaps:
            p = doc.add_paragraph()
            severity = gap.get('severity', 'medium')
            
            severity_run = p.add_run(f"[{severity.upper()}] ")
            severity_run.bold = True
            if severity == 'critical':
                severity_run.font.color.rgb = RGBColor(229, 62, 62)
            elif severity == 'high':
                severity_run.font.color.rgb = RGBColor(221, 107, 32)
            
            p.add_run(gap.get('gap', ''))
            
            if gap.get('suggestion'):
                suggestion = doc.add_paragraph()
                suggestion.add_run('Suggestion: ').italic = True
                suggestion.add_run(gap.get('suggestion', '')).italic = True
    
    # Footer
    doc.add_paragraph()
    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_run = footer.add_run('Generated by BRD Agent | Powered by Gemini AI')
    footer_run.font.color.rgb = RGBColor(160, 174, 192)
    footer_run.font.size = Pt(9)
    
    # Save to bytes
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    
    if output_path:
        with open(output_path, 'wb') as f:
            f.write(docx_bytes)
    
    return docx_bytes


if __name__ == "__main__":
    print(f"python-docx available: {_check_docx()}")
    
    if _check_docx():
        test_data = {
            "project_name": "Test BRD Export",
            "functional_reqs": [
                {"id": "FR-001", "text": "User authentication via SSO", "priority": 1, "moscow": "Must", "confidence": 0.95, "reasoning": "Explicit requirement with 'must' keyword"},
            ],
            "nfrs": [
                {"id": "NFR-001", "text": "Response time under 200ms", "category": "Performance", "confidence": 0.88},
            ],
            "stakeholders": [
                {"name": "John Smith", "role": "Product Manager", "influence_score": 0.9, "mentioned_count": 5},
            ],
            "decisions": [
                {"id": "DEC-001", "text": "Use PostgreSQL for database", "rationale": "Better JSON support"},
            ],
            "completeness_score": 0.85,
        }
        
        docx_bytes = export_to_word(test_data)
        print(f"Generated DOCX: {len(docx_bytes)} bytes")
